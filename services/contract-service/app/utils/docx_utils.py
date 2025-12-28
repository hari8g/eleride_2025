"""
Utilities for DOCX manipulation and placeholder replacement.
Handles placeholders even when Word splits text across runs.
"""
from typing import Dict, List
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def replace_placeholder_in_paragraph(paragraph, placeholder: str, value: str) -> bool:
    """
    Replace placeholder in a paragraph, handling Word's text runs.
    Returns True if replacement occurred.
    """
    if placeholder not in paragraph.text:
        return False
    
    # Get all text runs
    runs = paragraph.runs
    if not runs:
        return False
    
    # Rebuild paragraph text to find placeholder
    full_text = paragraph.text
    if placeholder not in full_text:
        return False
    
    # Replace in full text
    new_text = full_text.replace(placeholder, str(value))
    
    # Clear existing runs
    paragraph.clear()
    
    # Add new text
    paragraph.add_run(new_text)
    
    return True


def replace_placeholder_in_table(table, placeholder: str, value: str) -> int:
    """
    Replace placeholder in all cells of a table.
    Returns count of replacements.
    """
    count = 0
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if replace_placeholder_in_paragraph(paragraph, placeholder, value):
                    count += 1
    return count


def replace_placeholders_in_document(doc: Document, placeholders: Dict[str, str]) -> int:
    """
    Replace all placeholders in document (paragraphs, tables, headers, footers).
    Returns total count of replacements.
    """
    total_replacements = 0
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if replace_placeholder_in_paragraph(paragraph, placeholder, value):
                total_replacements += 1
    
    # Replace in tables
    for table in doc.tables:
        for placeholder, value in placeholders.items():
            count = replace_placeholder_in_table(table, placeholder, value)
            total_replacements += count
    
    # Replace in headers and footers
    for section in doc.sections:
        # Header
        for paragraph in section.header.paragraphs:
            for placeholder, value in placeholders.items():
                if replace_placeholder_in_paragraph(paragraph, placeholder, value):
                    total_replacements += 1
        
        # Footer
        for paragraph in section.footer.paragraphs:
            for placeholder, value in placeholders.items():
                if replace_placeholder_in_paragraph(paragraph, placeholder, value):
                    total_replacements += 1
    
    return total_replacements


def find_placeholders_in_document(doc: Document) -> List[str]:
    """
    Find all placeholders in the document (format: {{PLACEHOLDER}}).
    Returns list of unique placeholders found.
    """
    placeholders = set()
    
    def extract_placeholders(text: str):
        import re
        pattern = r'\{\{([A-Z_][A-Z0-9_]*)\}\}'
        matches = re.findall(pattern, text)
        placeholders.update(matches)
    
    # Check paragraphs
    for paragraph in doc.paragraphs:
        extract_placeholders(paragraph.text)
    
    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    extract_placeholders(paragraph.text)
    
    # Check headers and footers
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            extract_placeholders(paragraph.text)
        for paragraph in section.footer.paragraphs:
            extract_placeholders(paragraph.text)
    
    return sorted(list(placeholders))

